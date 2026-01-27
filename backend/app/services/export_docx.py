"""DOCX export service for manuscripts.

Uses python-docx to generate properly formatted Word documents
with industry-standard manuscript formatting for traditional publisher submissions.

Manuscript standards implemented:
- Font: Times New Roman, 12pt
- Line spacing: 2.0 (double-spaced)
- First line indent: 0.5 inches
- Margins: 1 inch all sides
- New page per chapter
"""

import io
import logging
import re
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.services.manuscript_assembly import (
    ChapterContent,
    CopyrightContent,
    DedicationContent,
    EpigraphContent,
    Manuscript,
    TitlePageContent,
)

from .exporters.base import BaseExporter, ExporterRegistry, ExportFormat, ExportResult

logger = logging.getLogger(__name__)


class DOCXExportError(Exception):
    """Exception raised when DOCX export fails."""

    pass


class DOCXExportService(BaseExporter):
    """Service for exporting manuscripts to DOCX format.

    Features:
    - Industry-standard manuscript formatting
    - Times New Roman 12pt font
    - Double-spaced (2.0 line spacing)
    - 0.5 inch first-line indent
    - 1 inch margins on all sides
    - Front matter sections (title page, copyright, dedication)
    - Back matter sections (author bio, also by, excerpt)
    - Chapter formatting with page breaks
    - Scene break handling
    - Core document properties (title, author)
    """

    format = ExportFormat.DOCX
    file_extension = ".docx"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # Manuscript formatting standards
    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(12)
    LINE_SPACING = 2.0  # Double-spaced
    FIRST_LINE_INDENT = Inches(0.5)
    MARGIN_TOP = Inches(1)
    MARGIN_BOTTOM = Inches(1)
    MARGIN_LEFT = Inches(1)
    MARGIN_RIGHT = Inches(1)

    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 12,
        line_spacing: float = 2.0,
        first_line_indent: float = 0.5,
    ):
        """Initialize the DOCX export service.

        Args:
            font_name: Font to use (defaults to Times New Roman).
            font_size: Font size in points (defaults to 12).
            line_spacing: Line spacing multiplier (defaults to 2.0 for double-spacing).
            first_line_indent: First line indent in inches (defaults to 0.5).
        """
        self.font_name = font_name
        self.font_size = Pt(font_size)
        self.line_spacing = line_spacing
        self.first_line_indent = Inches(first_line_indent)

    def export(self, manuscript: Manuscript) -> ExportResult:
        """Export the manuscript as a DOCX file.

        Args:
            manuscript: The manuscript to export.

        Returns:
            ExportResult with the DOCX content and metadata.
        """
        try:
            content = self._generate_docx(manuscript)
            return self._success_result(content, manuscript)
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return self._error_result(f"DOCX export failed: {str(e)}", manuscript)

    def _generate_docx(self, manuscript: Manuscript) -> bytes:
        """Generate the DOCX document.

        Args:
            manuscript: The manuscript to convert.

        Returns:
            Bytes content of the DOCX file.
        """
        doc = Document()

        # Set up document properties
        self._set_document_properties(doc, manuscript)

        # Set up page layout
        self._set_page_layout(doc)

        # Set up styles
        self._set_up_styles(doc)

        # Build document content
        self._build_document(doc, manuscript)

        # Write to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _set_document_properties(self, doc: Document, manuscript: Manuscript) -> None:
        """Set document core properties.

        Args:
            doc: The Document object.
            manuscript: The manuscript for metadata.
        """
        core_props = doc.core_properties
        core_props.title = manuscript.title
        if manuscript.author_name:
            core_props.author = manuscript.author_name

    def _set_page_layout(self, doc: Document) -> None:
        """Set up page layout with margins.

        Args:
            doc: The Document object.
        """
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT

    def _set_up_styles(self, doc: Document) -> None:
        """Set up document styles for manuscript formatting.

        Args:
            doc: The Document object.
        """
        styles = doc.styles

        # Normal style (body text)
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = self.font_name
        normal_font.size = self.font_size
        normal_pf = normal_style.paragraph_format
        normal_pf.line_spacing = self.line_spacing
        normal_pf.first_line_indent = self.first_line_indent
        normal_pf.space_after = Pt(0)
        normal_pf.space_before = Pt(0)

        # Title style
        self._create_style(
            styles,
            "ManuscriptTitle",
            base_style="Normal",
            font_size=Pt(24),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(12),
            first_line_indent=Inches(0),
        )

        # Subtitle style
        self._create_style(
            styles,
            "ManuscriptSubtitle",
            base_style="Normal",
            font_size=Pt(16),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(6),
            first_line_indent=Inches(0),
        )

        # Author style
        self._create_style(
            styles,
            "ManuscriptAuthor",
            base_style="Normal",
            font_size=Pt(14),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(24),
            first_line_indent=Inches(0),
        )

        # Chapter heading style
        self._create_style(
            styles,
            "ChapterHeading",
            base_style="Normal",
            font_size=Pt(14),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(72),
            space_after=Pt(6),
            first_line_indent=Inches(0),
        )

        # Chapter title style
        self._create_style(
            styles,
            "ChapterTitle",
            base_style="Normal",
            font_size=Pt(12),
            italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(36),
            first_line_indent=Inches(0),
        )

        # Section heading style
        self._create_style(
            styles,
            "SectionHeading",
            base_style="Normal",
            font_size=Pt(14),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(24),
            space_after=Pt(12),
            first_line_indent=Inches(0),
        )

        # Copyright style
        self._create_style(
            styles,
            "Copyright",
            base_style="Normal",
            font_size=Pt(10),
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=Pt(8),
            first_line_indent=Inches(0),
        )

        # Dedication style
        self._create_style(
            styles,
            "Dedication",
            base_style="Normal",
            italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=Inches(0),
        )

        # Epigraph style
        self._create_style(
            styles,
            "Epigraph",
            base_style="Normal",
            italic=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            left_indent=Inches(1),
            right_indent=Inches(1),
            first_line_indent=Inches(0),
        )

        # Attribution style
        self._create_style(
            styles,
            "Attribution",
            base_style="Normal",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent=Inches(1),
            space_before=Pt(12),
            first_line_indent=Inches(0),
        )

        # Scene break style
        self._create_style(
            styles,
            "SceneBreak",
            base_style="Normal",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(18),
            space_after=Pt(18),
            first_line_indent=Inches(0),
        )

        # First paragraph style (no indent)
        self._create_style(
            styles,
            "FirstParagraph",
            base_style="Normal",
            first_line_indent=Inches(0),
        )

    def _create_style(
        self,
        styles,
        name: str,
        base_style: str = "Normal",
        font_size: Optional[Pt] = None,
        bold: bool = False,
        italic: bool = False,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
        space_before: Optional[Pt] = None,
        space_after: Optional[Pt] = None,
        first_line_indent: Optional[Inches] = None,
        left_indent: Optional[Inches] = None,
        right_indent: Optional[Inches] = None,
    ) -> None:
        """Create or update a paragraph style.

        Args:
            styles: Document styles collection.
            name: Name for the style.
            base_style: Name of the base style to inherit from.
            font_size: Font size in points.
            bold: Whether to make text bold.
            italic: Whether to make text italic.
            alignment: Paragraph alignment.
            space_before: Space before paragraph.
            space_after: Space after paragraph.
            first_line_indent: First line indent.
            left_indent: Left indent for entire paragraph.
            right_indent: Right indent for entire paragraph.
        """
        # Check if style exists
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[base_style]

        # Font settings
        font = style.font
        font.name = self.font_name
        if font_size:
            font.size = font_size
        else:
            font.size = self.font_size
        font.bold = bold
        font.italic = italic

        # Paragraph settings
        pf = style.paragraph_format
        pf.line_spacing = self.line_spacing

        if alignment is not None:
            pf.alignment = alignment
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        if left_indent is not None:
            pf.left_indent = left_indent
        if right_indent is not None:
            pf.right_indent = right_indent

    def _build_document(self, doc: Document, manuscript: Manuscript) -> None:
        """Build the complete document content.

        Args:
            doc: The Document object.
            manuscript: The manuscript to convert.
        """
        # Front matter
        if manuscript.title_page:
            self._add_title_page(doc, manuscript.title_page)
            self._add_page_break(doc)

        if manuscript.copyright_page:
            self._add_copyright_page(doc, manuscript.copyright_page)
            self._add_page_break(doc)

        if manuscript.dedication:
            self._add_dedication_page(doc, manuscript.dedication)
            self._add_page_break(doc)

        if manuscript.epigraph:
            self._add_epigraph_page(doc, manuscript.epigraph)
            self._add_page_break(doc)

        if manuscript.acknowledgments:
            self._add_acknowledgments_page(doc, manuscript.acknowledgments)
            self._add_page_break(doc)

        # Chapters
        for i, chapter in enumerate(manuscript.chapters):
            if not chapter.content or not chapter.content.strip():
                continue  # Skip empty chapters

            self._add_chapter(doc, chapter)

            # Page break after each chapter except the last
            if i < len(manuscript.chapters) - 1:
                self._add_page_break(doc)

        # Back matter
        if manuscript.author_bio:
            self._add_page_break(doc)
            self._add_author_bio_page(doc, manuscript)

        if manuscript.also_by:
            self._add_page_break(doc)
            self._add_also_by_page(doc, manuscript)

        if manuscript.excerpt:
            self._add_page_break(doc)
            self._add_excerpt_page(doc, manuscript)

    def _add_page_break(self, doc: Document) -> None:
        """Add a page break to the document.

        Args:
            doc: The Document object.
        """
        doc.add_page_break()

    def _add_title_page(self, doc: Document, content: TitlePageContent) -> None:
        """Add the title page.

        Args:
            doc: The Document object.
            content: Title page content.
        """
        # Add vertical spacing
        for _ in range(6):
            doc.add_paragraph("", style="Normal")

        # Title
        doc.add_paragraph(content.title, style="ManuscriptTitle")

        # Subtitle
        if content.subtitle:
            doc.add_paragraph(content.subtitle, style="ManuscriptSubtitle")

        # Author
        if content.author_name:
            doc.add_paragraph(f"by {content.author_name}", style="ManuscriptAuthor")

        # Publisher
        if content.publisher:
            for _ in range(3):
                doc.add_paragraph("", style="Normal")
            doc.add_paragraph(content.publisher, style="ManuscriptAuthor")

        # Edition
        if content.edition:
            doc.add_paragraph(content.edition, style="ManuscriptAuthor")

    def _add_copyright_page(self, doc: Document, content: CopyrightContent) -> None:
        """Add the copyright page.

        Args:
            doc: The Document object.
            content: Copyright page content.
        """
        # Add vertical spacing
        for _ in range(6):
            doc.add_paragraph("", style="Normal")

        # Copyright notice
        author = content.author_name or "Author"
        doc.add_paragraph(f"Copyright \u00a9 {content.year} {author}", style="Copyright")

        # Rights statement
        doc.add_paragraph(content.rights_statement, style="Copyright")

        # Publisher
        if content.publisher:
            doc.add_paragraph("", style="Normal")
            doc.add_paragraph(f"Published by {content.publisher}", style="Copyright")

        # ISBN
        if content.isbn:
            doc.add_paragraph(f"ISBN: {content.isbn}", style="Copyright")

        # Edition info
        if content.edition_info:
            doc.add_paragraph("", style="Normal")
            doc.add_paragraph(content.edition_info, style="Copyright")

        # Credits
        if content.credits:
            doc.add_paragraph("", style="Normal")
            for credit in content.credits:
                doc.add_paragraph(credit, style="Copyright")

        # Fiction disclaimer
        doc.add_paragraph("", style="Normal")
        doc.add_paragraph("", style="Normal")
        disclaimer = (
            "This is a work of fiction. Names, characters, places, and incidents "
            "either are the product of the author's imagination or are used "
            "fictitiously. Any resemblance to actual persons, living or dead, "
            "events, or locales is entirely coincidental."
        )
        doc.add_paragraph(disclaimer, style="Copyright")

    def _add_dedication_page(self, doc: Document, content: DedicationContent) -> None:
        """Add the dedication page.

        Args:
            doc: The Document object.
            content: Dedication content.
        """
        # Add vertical spacing
        for _ in range(8):
            doc.add_paragraph("", style="Normal")

        # Dedication text
        doc.add_paragraph(content.text, style="Dedication")

    def _add_epigraph_page(self, doc: Document, content: EpigraphContent) -> None:
        """Add the epigraph page.

        Args:
            doc: The Document object.
            content: Epigraph content.
        """
        # Add vertical spacing
        for _ in range(8):
            doc.add_paragraph("", style="Normal")

        # Quote
        doc.add_paragraph(f'"{content.text}"', style="Epigraph")

        # Attribution
        if content.attribution or content.source:
            attr_parts = []
            if content.attribution:
                attr_parts.append(content.attribution)
            if content.source:
                attr_parts.append(content.source)
            attribution = ", ".join(attr_parts)
            doc.add_paragraph(f"\u2014 {attribution}", style="Attribution")

    def _add_acknowledgments_page(self, doc: Document, text: str) -> None:
        """Add the acknowledgments page.

        Args:
            doc: The Document object.
            text: Acknowledgments text.
        """
        # Heading
        doc.add_paragraph("Acknowledgments", style="SectionHeading")

        # Content
        paragraphs = self._text_to_paragraphs(text)
        for i, para in enumerate(paragraphs):
            style = "FirstParagraph" if i == 0 else "Normal"
            doc.add_paragraph(para, style=style)

    def _add_chapter(self, doc: Document, chapter: ChapterContent) -> None:
        """Add a chapter.

        Args:
            doc: The Document object.
            chapter: Chapter content.
        """
        # Chapter heading
        doc.add_paragraph(f"CHAPTER {chapter.number}", style="ChapterHeading")

        # Chapter title
        if chapter.title:
            doc.add_paragraph(chapter.title, style="ChapterTitle")

        # Chapter content
        self._add_content(doc, chapter.content)

    def _add_content(self, doc: Document, content: str) -> None:
        """Add chapter content with proper formatting.

        Handles:
        - Paragraph breaks (double newlines)
        - Scene breaks (* * *, # # #, etc.)
        - First paragraph styling (no indent)

        Args:
            doc: The Document object.
            content: Raw chapter content text.
        """
        if not content:
            return

        # Normalize line endings
        content = content.replace("\r\n", "\n").replace("\r", "\n")

        # Handle scene breaks
        scene_break_patterns = [
            r"\n\s*\*\s*\*\s*\*\s*\n",
            r"\n\s*#\s*#\s*#\s*\n",
            r"\n\s*-\s*-\s*-\s*\n",
            r"\n\s*~\s*~\s*~\s*\n",
        ]
        scene_break_marker = "\n\n[SCENE_BREAK]\n\n"
        for pattern in scene_break_patterns:
            content = re.sub(pattern, scene_break_marker, content)

        # Split into paragraphs
        paragraphs = re.split(r"\n\s*\n", content)

        first_after_break = True

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if para == "[SCENE_BREAK]":
                doc.add_paragraph("* * *", style="SceneBreak")
                first_after_break = True
            else:
                style = "FirstParagraph" if first_after_break else "Normal"
                doc.add_paragraph(para, style=style)
                first_after_break = False

    def _add_author_bio_page(self, doc: Document, manuscript: Manuscript) -> None:
        """Add the author bio page.

        Args:
            doc: The Document object.
            manuscript: The manuscript with author bio.
        """
        bio = manuscript.author_bio

        # Heading
        doc.add_paragraph("About the Author", style="SectionHeading")

        # Bio text
        paragraphs = self._text_to_paragraphs(bio.text)
        for i, para in enumerate(paragraphs):
            style = "FirstParagraph" if i == 0 else "Normal"
            doc.add_paragraph(para, style=style)

        # Website
        if bio.website:
            doc.add_paragraph("", style="Normal")
            doc.add_paragraph(f"Website: {bio.website}", style="FirstParagraph")

        # Social media
        if bio.social_media:
            for platform, handle in bio.social_media.items():
                doc.add_paragraph(f"{platform.capitalize()}: {handle}", style="FirstParagraph")

    def _add_also_by_page(self, doc: Document, manuscript: Manuscript) -> None:
        """Add the "Also By" page.

        Args:
            doc: The Document object.
            manuscript: The manuscript with also_by content.
        """
        ab = manuscript.also_by

        # Heading
        doc.add_paragraph(f"Also by {ab.author_name}", style="SectionHeading")

        # Series info
        if ab.series_info:
            for series_name, titles in ab.series_info.items():
                doc.add_paragraph("", style="Normal")
                p = doc.add_paragraph(style="FirstParagraph")
                run = p.add_run(f"{series_name} Series")
                run.bold = True

                for title in titles:
                    p = doc.add_paragraph(style="FirstParagraph")
                    run = p.add_run(f"    {title}")
                    run.italic = True

        # Standalone titles
        if ab.titles:
            doc.add_paragraph("", style="Normal")
            p = doc.add_paragraph(style="FirstParagraph")
            run = p.add_run("Standalone Novels")
            run.bold = True

            for title in ab.titles:
                p = doc.add_paragraph(style="FirstParagraph")
                run = p.add_run(f"    {title}")
                run.italic = True

    def _add_excerpt_page(self, doc: Document, manuscript: Manuscript) -> None:
        """Add the excerpt/preview page.

        Args:
            doc: The Document object.
            manuscript: The manuscript with excerpt content.
        """
        ex = manuscript.excerpt

        # Coming soon
        if ex.coming_soon_date:
            p = doc.add_paragraph(style="FirstParagraph")
            run = p.add_run(f"Coming Soon: {ex.coming_soon_date}")
            run.bold = True
            doc.add_paragraph("", style="Normal")

        # Book title
        doc.add_paragraph(ex.book_title, style="SectionHeading")

        # Chapter title
        if ex.chapter_title:
            doc.add_paragraph(ex.chapter_title, style="ChapterTitle")

        # Excerpt text
        paragraphs = self._text_to_paragraphs(ex.text)
        for i, para in enumerate(paragraphs):
            style = "FirstParagraph" if i == 0 else "Normal"
            doc.add_paragraph(para, style=style)

    def _text_to_paragraphs(self, text: str) -> list[str]:
        """Convert plain text to list of paragraph strings.

        Args:
            text: Plain text content.

        Returns:
            List of paragraph strings.
        """
        if not text:
            return []

        paragraphs = text.strip().split("\n\n")
        return [p.strip() for p in paragraphs if p.strip()]


# Register the exporter
ExporterRegistry.register(ExportFormat.DOCX, DOCXExportService)
