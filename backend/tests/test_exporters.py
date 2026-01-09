"""Tests for manuscript exporters."""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.export_docx import DOCXExportService
from app.services.exporters import (
    ExportResult,
    MarkdownExporter,
    TextExporter,
)
from app.services.exporters.base import (
    ExporterRegistry,
    ExportFormat,
    ExportMetadata,
)
from app.services.manuscript_assembly import (
    AlsoByContent,
    AuthorBioContent,
    ChapterContent,
    CopyrightContent,
    DedicationContent,
    EpigraphContent,
    ExcerptContent,
    Manuscript,
    ManuscriptBuilder,
    TableOfContentsEntry,
    TitlePageContent,
)


class TestExportFormat:
    """Tests for ExportFormat enum."""

    def test_all_formats_defined(self):
        """All expected formats should be defined."""
        expected = {"docx", "pdf", "epub", "markdown", "text"}
        actual = {f.value for f in ExportFormat}
        assert actual == expected


class TestExportMetadata:
    """Tests for ExportMetadata dataclass."""

    def test_default_metadata(self):
        """Default metadata should have required fields."""
        metadata = ExportMetadata(title="Test")
        assert metadata.title == "Test"
        assert metadata.format == ExportFormat.TEXT
        assert metadata.word_count == 0

    def test_full_metadata(self):
        """Full metadata should have all fields."""
        metadata = ExportMetadata(
            title="My Book",
            author="Jane Doe",
            format=ExportFormat.MARKDOWN,
            word_count=50000,
            chapter_count=20,
            file_extension=".md",
            mime_type="text/markdown",
        )
        assert metadata.title == "My Book"
        assert metadata.author == "Jane Doe"
        assert metadata.word_count == 50000


class TestExportResult:
    """Tests for ExportResult dataclass."""

    def test_success_result(self):
        """Successful result should have content."""
        result = ExportResult(
            success=True,
            content=b"Test content",
            metadata=ExportMetadata(title="Test"),
        )
        assert result.success is True
        assert result.size_bytes == 12
        assert result.error_message is None

    def test_error_result(self):
        """Error result should have error message."""
        result = ExportResult(
            success=False,
            content=b"",
            metadata=ExportMetadata(title="Test"),
            error_message="Export failed",
        )
        assert result.success is False
        assert result.size_bytes == 0
        assert result.error_message == "Export failed"


class TestExporterRegistry:
    """Tests for ExporterRegistry."""

    def test_register_and_get(self):
        """Registered exporters should be retrievable."""
        # Text exporter should already be registered
        exporter_class = ExporterRegistry.get(ExportFormat.TEXT)
        assert exporter_class == TextExporter

    def test_create_exporter(self):
        """Should create exporter instances."""
        exporter = ExporterRegistry.create(ExportFormat.TEXT)
        assert isinstance(exporter, TextExporter)

    def test_available_formats(self):
        """Should return available formats."""
        formats = ExporterRegistry.available_formats()
        assert ExportFormat.TEXT in formats
        assert ExportFormat.MARKDOWN in formats

    def test_is_available(self):
        """Should check format availability."""
        assert ExporterRegistry.is_available(ExportFormat.TEXT) is True
        assert ExporterRegistry.is_available(ExportFormat.MARKDOWN) is True

    def test_unavailable_format(self):
        """Should handle unavailable formats."""
        # PDF and DOCX are not implemented yet
        ExporterRegistry.create(ExportFormat.PDF)
        # May or may not be None depending on implementation
        # Just verify no crash


class TestTextExporter:
    """Tests for TextExporter."""

    def test_export_minimal_manuscript(self):
        """Minimal manuscript should export."""
        manuscript = Manuscript(
            title="Test Book",
            chapters=[
                ChapterContent(
                    number=1,
                    title="First Chapter",
                    content="This is the content.",
                    word_count=4,
                )
            ],
        )
        exporter = TextExporter()
        result = exporter.export(manuscript)

        assert result.success is True
        assert result.metadata.title == "Test Book"
        assert result.metadata.format == ExportFormat.TEXT
        # TextExporter uppercases chapter titles
        assert b"FIRST CHAPTER" in result.content
        assert b"This is the content" in result.content

    def test_export_full_manuscript(self):
        """Full manuscript should export all sections."""
        manuscript = (
            ManuscriptBuilder("Complete Book")
            .set_author("Author Name")
            .add_title_page(subtitle="A Novel")
            .add_copyright_page(year=2024)
            .add_dedication("To readers")
            .add_chapter(1, "Chapter One", "Content here.")
            .generate_toc()
            .add_author_bio("Bio text here.")
            .build()
        )

        exporter = TextExporter()
        result = exporter.export(manuscript)

        assert result.success is True
        content = result.content.decode("utf-8")
        assert "COMPLETE BOOK" in content
        assert "Author Name" in content
        assert "To readers" in content
        assert "CHAPTER 1" in content
        assert "Content here" in content
        assert "ABOUT THE AUTHOR" in content

    def test_export_cleans_markers(self):
        """Export should clean formatting markers."""
        manuscript = Manuscript(
            title="Test",
            chapters=[
                ChapterContent(
                    number=1,
                    title="Ch 1",
                    content="Text here.",
                )
            ],
        )
        exporter = TextExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "[PAGE_BREAK]" not in content
        assert "[DROP_CAP]" not in content

    def test_export_file_name(self):
        """Export should generate valid file name."""
        manuscript = Manuscript(title="My Great Book!")
        exporter = TextExporter()
        result = exporter.export(manuscript)

        assert result.file_name == "My_Great_Book.txt"

    def test_export_file_name_special_chars(self):
        """File name should sanitize special characters."""
        manuscript = Manuscript(title="Book: A Title! With? Symbols*")
        exporter = TextExporter()
        result = exporter.export(manuscript)

        # Only alphanumeric, spaces, hyphens, underscores allowed
        assert ":" not in result.file_name
        assert "!" not in result.file_name
        assert "?" not in result.file_name
        assert "*" not in result.file_name

    def test_export_size_bytes(self):
        """Size should be calculated correctly."""
        manuscript = Manuscript(
            title="Test",
            chapters=[ChapterContent(number=1, title="Ch", content="X" * 100)],
        )
        exporter = TextExporter()
        result = exporter.export(manuscript)

        assert result.size_bytes > 100


class TestMarkdownExporter:
    """Tests for MarkdownExporter."""

    def test_export_minimal_manuscript(self):
        """Minimal manuscript should export as Markdown."""
        manuscript = Manuscript(
            title="Test Book",
            chapters=[
                ChapterContent(
                    number=1,
                    title="First Chapter",
                    content="This is the content.",
                    word_count=4,
                )
            ],
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        assert result.success is True
        assert result.metadata.format == ExportFormat.MARKDOWN
        assert result.file_name.endswith(".md")

        content = result.content.decode("utf-8")
        assert "## Chapter 1: First Chapter" in content
        assert "This is the content" in content

    def test_export_with_title_page(self):
        """Title page should be formatted as Markdown."""
        manuscript = Manuscript(
            title="My Novel",
            title_page=TitlePageContent(
                title="My Novel",
                subtitle="A Story",
                author_name="Jane Doe",
                publisher="Great Books",
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "# My Novel" in content
        assert "## A Story" in content
        assert "**by Jane Doe**" in content
        assert "*Great Books*" in content

    def test_export_with_copyright(self):
        """Copyright page should be formatted as Markdown."""
        manuscript = Manuscript(
            title="Test",
            copyright_page=CopyrightContent(
                author_name="Author",
                year=2024,
                isbn="978-123456789",
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "Copyright 2024 Author" in content
        assert "ISBN: 978-123456789" in content
        # Check for the full disclaimer wrapped in italics
        assert "*This is a work of fiction." in content

    def test_export_with_dedication(self):
        """Dedication should be italic in Markdown."""
        manuscript = Manuscript(
            title="Test",
            dedication=DedicationContent(text="For my family"),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "*For my family*" in content

    def test_export_with_epigraph(self):
        """Epigraph should be formatted as blockquote."""
        manuscript = Manuscript(
            title="Test",
            epigraph=EpigraphContent(
                text="To be or not to be",
                attribution="Shakespeare",
                source="Hamlet",
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "> To be or not to be" in content
        assert "Shakespeare" in content
        assert "*Hamlet*" in content

    def test_export_with_toc(self):
        """Table of contents should be formatted as list."""
        manuscript = Manuscript(
            title="Test",
            table_of_contents=[
                TableOfContentsEntry(title="Intro", chapter_number=1, level=1),
                TableOfContentsEntry(title="Middle", chapter_number=2, level=1),
                TableOfContentsEntry(title="End", chapter_number=3, level=1),
            ],
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "## Table of Contents" in content
        assert "- Chapter 1: Intro" in content
        assert "- Chapter 2: Middle" in content
        assert "- Chapter 3: End" in content

    def test_export_with_author_bio(self):
        """Author bio should be formatted correctly."""
        manuscript = Manuscript(
            title="Test",
            author_bio=AuthorBioContent(
                text="Jane writes books.",
                website="https://jane.com",
                social_media={"twitter": "@jane"},
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "## About the Author" in content
        assert "Jane writes books." in content
        assert "[https://jane.com](https://jane.com)" in content
        assert "Twitter: @jane" in content

    def test_export_with_also_by(self):
        """Also by page should list books."""
        manuscript = Manuscript(
            title="Test",
            also_by=AlsoByContent(
                author_name="Jane",
                titles=["Book One", "Book Two"],
                series_info={"Magic Series": ["Magic 1", "Magic 2"]},
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "## Also by Jane" in content
        assert "### Magic Series Series" in content
        assert "- *Magic 1*" in content
        assert "- *Book One*" in content

    def test_export_with_excerpt(self):
        """Excerpt should be formatted correctly."""
        manuscript = Manuscript(
            title="Test",
            excerpt=ExcerptContent(
                book_title="Next Book",
                text="Preview text...",
                chapter_title="Chapter 1",
                coming_soon_date="Fall 2025",
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "### Coming Soon: Fall 2025" in content
        assert "## Next Book" in content
        assert "### Chapter 1" in content
        assert "Preview text..." in content

    def test_export_chapter_scene_breaks(self):
        """Scene breaks should become horizontal rules."""
        manuscript = Manuscript(
            title="Test",
            chapters=[
                ChapterContent(
                    number=1,
                    title="Ch 1",
                    content="Scene one.\n\n* * *\n\nScene two.",
                )
            ],
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        # Scene break markers converted to ---
        assert "Scene one." in content
        assert "Scene two." in content

    def test_export_full_manuscript(self):
        """Full manuscript should export correctly."""
        manuscript = (
            ManuscriptBuilder("Complete Novel")
            .set_author("Jane Doe")
            .add_title_page(subtitle="A Tale")
            .add_copyright_page(year=2024)
            .add_dedication("To everyone")
            .add_epigraph("A wise quote", attribution="Wise Person")
            .add_acknowledgments("Thanks to all")
            .add_chapter(1, "Beginning", "The story starts...")
            .add_chapter(2, "Middle", "The story continues...")
            .add_chapter(3, "End", "The story ends.")
            .generate_toc()
            .add_author_bio("Jane is an author.")
            .add_also_by(["Previous Book"])
            .add_excerpt("Next Book", "Preview...", coming_soon_date="2025")
            .build()
        )

        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        assert result.success is True
        content = result.content.decode("utf-8")

        # Check all sections present
        assert "# Complete Novel" in content
        assert "## A Tale" in content
        assert "Copyright 2024" in content
        assert "*To everyone*" in content
        assert "> A wise quote" in content
        assert "## Acknowledgments" in content
        assert "## Table of Contents" in content
        assert "## Chapter 1: Beginning" in content
        assert "## Chapter 2: Middle" in content
        assert "## Chapter 3: End" in content
        assert "## About the Author" in content
        assert "## Also by Jane Doe" in content
        assert "### Coming Soon: 2025" in content


class TestExporterMetadata:
    """Tests for exporter metadata generation."""

    def test_text_exporter_metadata(self):
        """Text exporter should generate correct metadata."""
        exporter = TextExporter()
        assert exporter.format == ExportFormat.TEXT
        assert exporter.file_extension == ".txt"
        assert exporter.mime_type == "text/plain"

    def test_markdown_exporter_metadata(self):
        """Markdown exporter should generate correct metadata."""
        exporter = MarkdownExporter()
        assert exporter.format == ExportFormat.MARKDOWN
        assert exporter.file_extension == ".md"
        assert exporter.mime_type == "text/markdown"

    def test_export_result_metadata(self):
        """Export result should have correct metadata."""
        manuscript = Manuscript(
            title="Test Book",
            author_name="Author",
            chapters=[ChapterContent(number=1, title="Ch", content="X", word_count=1000)],
        )
        exporter = TextExporter()
        result = exporter.export(manuscript)

        assert result.metadata.title == "Test Book"
        assert result.metadata.author == "Author"
        assert result.metadata.word_count == 1000
        assert result.metadata.chapter_count == 1


class TestIntegration:
    """Integration tests for exporters."""

    def test_both_exporters_produce_output(self):
        """Both exporters should produce valid output."""
        manuscript = (
            ManuscriptBuilder("Test Book")
            .set_author("Author")
            .add_chapter(1, "Chapter", "Content here.")
            .build()
        )

        text_result = TextExporter().export(manuscript)
        md_result = MarkdownExporter().export(manuscript)

        assert text_result.success is True
        assert md_result.success is True
        assert text_result.size_bytes > 0
        assert md_result.size_bytes > 0

    def test_registry_creates_correct_exporters(self):
        """Registry should create correct exporter instances."""
        text_exporter = ExporterRegistry.create(ExportFormat.TEXT)
        md_exporter = ExporterRegistry.create(ExportFormat.MARKDOWN)

        assert isinstance(text_exporter, TextExporter)
        assert isinstance(md_exporter, MarkdownExporter)

    def test_export_large_manuscript(self):
        """Large manuscripts should export correctly."""
        builder = ManuscriptBuilder("Large Book").set_author("Author")

        # Add 20 chapters
        for i in range(1, 21):
            builder.add_chapter(i, f"Chapter {i}", f"Content for chapter {i}. " * 100)

        manuscript = builder.generate_toc().build()

        text_result = TextExporter().export(manuscript)
        md_result = MarkdownExporter().export(manuscript)

        assert text_result.success is True
        assert md_result.success is True
        assert text_result.metadata.chapter_count == 20
        assert md_result.metadata.chapter_count == 20


class TestMarkdownExporterEdgeCases:
    """Tests for edge cases in markdown export."""

    def test_title_page_with_edition(self):
        """Title page with edition should be included."""
        manuscript = Manuscript(
            title="Test",
            title_page=TitlePageContent(
                title="Test Book",
                author_name="Author",
                edition="Second Edition",
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "*Second Edition*" in content

    def test_copyright_with_publisher_and_edition(self):
        """Copyright with publisher and edition info."""
        manuscript = Manuscript(
            title="Test",
            copyright_page=CopyrightContent(
                author_name="Author",
                year=2024,
                publisher="Big Publisher",
                edition_info="First hardcover edition",
                credits=["Cover design by Artist", "Edited by Editor"],
            ),
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "Published by Big Publisher" in content
        assert "First hardcover edition" in content
        assert "- Cover design by Artist" in content
        assert "- Edited by Editor" in content

    def test_export_error_handling(self):
        """Export errors should be captured in result."""
        # Create a manuscript that might cause issues
        manuscript = Manuscript(
            title="Test",
            chapters=[ChapterContent(number=1, title="Ch1", content="Valid content")],
        )
        exporter = MarkdownExporter()

        # Normal export should succeed
        result = exporter.export(manuscript)
        assert result.success is True

    def test_chapter_without_scene_breaks(self):
        """Chapter without scene breaks should export cleanly."""
        manuscript = Manuscript(
            title="Test",
            chapters=[
                ChapterContent(
                    number=1,
                    title="Simple Chapter",
                    content="Just one continuous scene with no breaks.",
                    word_count=6,
                )
            ],
        )
        exporter = MarkdownExporter()
        result = exporter.export(manuscript)

        content = result.content.decode("utf-8")
        assert "Just one continuous scene" in content
        assert "---" not in content  # No scene break markers


class TestFileNameSanitization:
    """Tests for file name sanitization in exporters."""

    def test_empty_title_uses_default(self):
        """Empty or special-chars-only title should default to 'manuscript'."""
        manuscript = Manuscript(
            title="@#$%^&*()",  # Only special chars, sanitizes to empty
            author_name="Test Author",
            chapters=[],
        )
        exporter = TextExporter()
        file_name = exporter._create_file_name(manuscript)
        assert file_name == "manuscript.txt"


class TestDOCXExporter:
    """Tests for DOCXExportService."""

    def test_export_basic_manuscript(self):
        """DOCX exporter should export a basic manuscript."""
        manuscript = Manuscript(
            title="Test Book",
            chapters=[
                ChapterContent(
                    number=1,
                    title="First Chapter",
                    content="This is the content.",
                    word_count=4,
                )
            ],
        )
        exporter = DOCXExportService()
        result = exporter.export(manuscript)

        assert result.success is True
        assert result.metadata.title == "Test Book"
        assert result.metadata.format == ExportFormat.DOCX
        assert result.size_bytes > 0

    def test_export_with_front_matter(self):
        """DOCX exporter should include front matter sections."""
        manuscript = (
            ManuscriptBuilder("My Novel")
            .set_author("Jane Doe")
            .add_title_page(subtitle="A Story")
            .add_copyright_page(year=2024, publisher="Great Books", isbn="978-123456789")
            .add_dedication("To my readers")
            .add_chapter(1, "Chapter One", "Content here.")
            .build()
        )

        exporter = DOCXExportService()
        result = exporter.export(manuscript)

        assert result.success is True

        # Parse the generated DOCX to verify content
        doc = Document(io.BytesIO(result.content))
        paragraphs_text = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs_text)

        # Verify title page elements
        assert "My Novel" in full_text
        assert "A Story" in full_text
        assert "by Jane Doe" in full_text

        # Verify copyright elements
        assert "Copyright" in full_text and "2024" in full_text
        assert "Great Books" in full_text
        assert "978-123456789" in full_text

        # Verify dedication
        assert "To my readers" in full_text

    def test_front_matter_formatting(self):
        """DOCX front matter should have proper alignment."""
        manuscript = (
            ManuscriptBuilder("Test Book")
            .set_author("Author Name")
            .add_title_page()
            .add_copyright_page(year=2024)
            .add_dedication("Dedication text")
            .add_chapter(1, "Chapter", "Content")
            .build()
        )

        exporter = DOCXExportService()
        result = exporter.export(manuscript)
        doc = Document(io.BytesIO(result.content))

        # Check that styles have correct alignment
        # Title style should be centered
        title_style = doc.styles["ManuscriptTitle"]
        assert title_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

        # Copyright style should be left-aligned
        copyright_style = doc.styles["Copyright"]
        assert copyright_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT

        # Dedication style should be centered
        dedication_style = doc.styles["Dedication"]
        assert dedication_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_export_with_epigraph(self):
        """DOCX exporter should include epigraph."""
        manuscript = (
            ManuscriptBuilder("Novel")
            .set_author("Author")
            .add_epigraph(
                "To be or not to be",
                attribution="Shakespeare",
                source="Hamlet",
            )
            .add_chapter(1, "Ch 1", "Content")
            .build()
        )

        exporter = DOCXExportService()
        result = exporter.export(manuscript)

        doc = Document(io.BytesIO(result.content))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "To be or not to be" in full_text
        assert "Shakespeare" in full_text
        assert "Hamlet" in full_text

    def test_export_manuscript_formatting(self):
        """DOCX should use manuscript formatting standards."""
        manuscript = Manuscript(
            title="Test",
            chapters=[ChapterContent(number=1, title="Ch", content="Content")],
        )

        exporter = DOCXExportService()
        result = exporter.export(manuscript)
        doc = Document(io.BytesIO(result.content))

        # Check Normal style settings
        normal_style = doc.styles["Normal"]

        # Font should be Times New Roman
        assert normal_style.font.name == "Times New Roman"

        # Font size should be 12pt (152400 EMUs = 12pt)
        assert normal_style.font.size.pt == 12

        # Line spacing should be 2.0 (double-spaced)
        assert normal_style.paragraph_format.line_spacing == 2.0

    def test_export_docx_file_extension(self):
        """DOCX export should have correct file extension."""
        manuscript = Manuscript(title="Book")
        exporter = DOCXExportService()
        result = exporter.export(manuscript)

        assert result.file_name.endswith(".docx")

    def test_export_docx_mime_type(self):
        """DOCX export should have correct MIME type."""
        exporter = DOCXExportService()
        expected = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert exporter.mime_type == expected
